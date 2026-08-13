"""不需要網路 / 憑證的單元測試。跑法：python tests/test_pure.py"""

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
try:
    sys.stdout.reconfigure(encoding="utf-8")
except Exception:
    pass

import datetime as dt
import json

from core.docs_client import _render, _u16
from core.sheets_client import _a1, normalize_ticker, parse_call_datetime

failures = []


def check(name, got, want):
    if got == want:
        print(f"  [OK] {name}")
    else:
        print(f"  [X]  {name}: got={got!r} want={want!r}")
        failures.append(name)


print("== 代號正規化（表上寫法 → 真正的美股代號）")
check("CITI → C", normalize_ticker("CITI"), "C")
check("'C (Citi)' → C", normalize_ticker("C (Citi)"), "C")
check("Unity → U", normalize_ticker("Unity"), "U")
check("空白容錯", normalize_ticker("  vst "), "VST")
check("一般代號不動", normalize_ticker("RKLB"), "RKLB")

print("\n== 日期時間解析（表上沒有年份，時間是台灣時間的『會議開始』時刻）")
d, s = parse_call_datetime("  07/14   週二", " 19:30", year=2026)
check("正常列", (d, s), (dt.datetime(2026, 7, 14, 19, 30), ""))
check("沒有日期 → 跳過", parse_call_datetime("", "19:30", year=2026)[1], "no_date")
check("沒有時間 → 跳過", parse_call_datetime("07/14 週二", "", year=2026)[1], "no_time")
check("凌晨場次", parse_call_datetime("11/12 週三", "06:00", year=2026)[0],
      dt.datetime(2026, 11, 12, 6, 0))
check("非法日期不炸", parse_call_datetime("02/30", "06:00", year=2026)[0], None)

print("\n== A1 欄名（狀態欄會加在最右邊，可能超過 Z）")
check("0 → A", _a1(0), "A")
check("25 → Z", _a1(25), "Z")
check("26 → AA", _a1(26), "AA")
check("51 → AZ", _a1(51), "AZ")

print("\n== UTF-16 index（Docs API 的 index 以 UTF-16 code unit 計）")
check("ASCII", _u16("abc"), 3)
check("中文各佔 1 單位", _u16("財報"), 2)
check("emoji 佔 2 單位", _u16("🚀"), 2)

print("\n== memo JSON → Docs 內容渲染（用真實的 VST 2Q26 資料）")
src = Path(__file__).resolve().parent.parent.parent / "MEMO" / "_work" / "vst_2026q2.json"
if src.exists():
    print("   （來源：vst_2026q2.json）")
    memo = json.loads(src.read_text(encoding="utf-8"))
    text, bolds, headings = _render(memo, "https://drive.google.com/file/d/TEST", "2026/08/07",
                                    title="VST 2026Q2")
    lines = text.split("\n")
    check("第 1 行 = 分頁名（不是 memo 裡的公司全名）", lines[0], "VST 2026Q2")
    check("第 2 行 = 日期正規化為 YYYYMMDD", lines[1], "20260807")
    check("第 3 行 = podcast link 標籤", lines[2], "podcast link：")
    check("第 4 行 = podcast 連結", lines[3], "https://drive.google.com/file/d/TEST")
    check("有重點摘要", "重點摘要" in text, True)
    check("章節全數渲染",
          all(s["heading"] in text for s in memo["sections"]), True)
    ok_bold = all(0 <= a < b <= _u16(text) for a, b in bolds)
    check("粗體區間都在合法範圍內", ok_bold, True)
    ok_head = all(0 <= a < b <= _u16(text) for a, b, _ in headings)
    check("段落樣式區間都在合法範圍內", ok_head, True)
    print(f"       渲染結果：{_u16(text)} 字元、{len(bolds)} 段粗體、{len(headings)} 個標題")
else:
    print(f"  [--] 找不到 {src}，跳過渲染測試")

print("\n== 季度標籤推導（財報日期 → 最近結束的曆季）")
from core.quarters import calendar_quarter_for, local_folder_label, tab_name, fy_warning
import datetime as _dt
check("VST 2026-08-07 → 2026Q2", tab_name("VST", _dt.date(2026, 8, 7)), "VST 2026Q2")
check("本機資料夾 2Q26", local_folder_label(_dt.date(2026, 8, 7)), "2Q26")
check("跨年：1 月發 → 前一年 Q4", calendar_quarter_for(_dt.date(2026, 1, 29)), (2025, 4))
check("4 月發 → Q1", calendar_quarter_for(_dt.date(2026, 4, 23)), (2026, 1))
check("11 月發 → Q3", calendar_quarter_for(_dt.date(2026, 11, 20)), (2026, 3))
check("NVDA 有非曆年警告", fy_warning("NVDA") is not None, True)
check("VST 無警告", fy_warning("VST"), None)

print("\n== 美東→台灣時間換算（含日光節約）")
from core.calendar_source import estimate_call_time_tw
check("8月 AMC → 隔日 05:00", estimate_call_time_tw(_dt.date(2026, 8, 13), "AMC"),
      _dt.datetime(2026, 8, 14, 5, 0))
check("8月 BMO → 當日 20:30", estimate_call_time_tw(_dt.date(2026, 8, 20), "BMO"),
      _dt.datetime(2026, 8, 20, 20, 30))
check("1月 AMC → 隔日 06:00（冬令）", estimate_call_time_tw(_dt.date(2026, 1, 29), "AMC"),
      _dt.datetime(2026, 1, 30, 6, 0))
check("未知場次 → None", estimate_call_time_tw(_dt.date(2026, 8, 13), "?"), None)

print("\n== 逐字稿標題判讀（全部取自 2026-08-12 真實抓到的標題）")
from core.transcripts import PREVIEW, STRONG, WEAK, _leading_token_ok, _relevant, name_keywords

check("公司名關鍵詞：Lumentum", name_keywords("Lumentum Holdings Inc."), ["Lumentum"])
check("公司名關鍵詞：Super Micro", name_keywords("Super Micro Computer, Inc."),
      ["Super", "Micro", "Computer"])
check("公司名關鍵詞：CoreWeave", name_keywords("CoreWeave, Inc."), ["CoreWeave"])


def judge(title, ticker, company):
    """完整判準：強訊號 + 指向本公司 + 不是預告。"""
    kw = name_keywords(company)
    return (bool(STRONG.search(title))
            and not PREVIEW.search(title)
            and _relevant(title, ticker, kw)
            and _leading_token_ok(title, ticker, kw))


LITE, SMCI, CRWV = "Lumentum Holdings Inc.", "Super Micro Computer, Inc.", "CoreWeave, Inc."
check("Yahoo 逐字稿（只寫代號變體）",
      judge("SUPER MICRO COMPUTER INC (SMCI.MX) Q4 FY2026 earnings call transcript",
            "SMCI", SMCI), True)
check("Investing 逐字稿（標題只寫公司名）",
      judge("Earnings call transcript: Lumentum tops Q4 2026 forecasts as AI demand lifts outlook",
            "LITE", LITE), True)
check("MarketBeat call highlights",
      judge("CoreWeave Q2 Earnings Call Highlights", "CRWV", CRWV), True)
check("別家公司的電話會議只是提到我們 → 不算",
      judge("BLZE Q2 Earnings Call Focuses on CoreWeave and AI Storage", "CRWV", CRWV), False)
check("財報預告 → 不算",
      judge("Super Micro Computer Earnings Are Imminent; These Most Accurate Analysts "
            "Revise Forecasts Ahead Of Earnings Call", "SMCI", SMCI), False)
check("純股價新聞 → 不算",
      judge("Super Micro stock jumps on upbeat outlook", "SMCI", SMCI), False)
check("別家公司的逐字稿 → 不算",
      judge("Earnings call transcript: TETRA Technologies tops revenue forecast in Q2 2026",
            "LITE", LITE), False)
check("SA 法說簡報只是弱訊號（不足以判定就緒）",
      bool(STRONG.search("CoreWeave, Inc. 2026 Q2 - Results - Earnings Call Presentation")), False)
check("SA 法說簡報仍列為線索",
      bool(WEAK.search("CoreWeave, Inc. 2026 Q2 - Results - Earnings Call Presentation")), True)

print("\n== 用量上限判讀（取自 2026-08-12 真實的 agent 輸出）")
sys.path.insert(0, str(Path(__file__).resolve().parent.parent / "tools"))
from dispatcher import QUOTA, RESET_AT

REAL = "You've hit your session limit · resets 10:20am (Asia/Taipei)\n"
check("真實的用量上限訊息判得出來", bool(QUOTA.search(REAL)), True)
check("解析得出恢復時間", RESET_AT.search(REAL).group(1).strip(), "resets 10:20am (Asia/Taipei)"[7:].strip())
check("一般的 agent 錯誤不會誤判",
      bool(QUOTA.search("Traceback: KeyError 'sections'")), False)
check("找不到檔案不會誤判",
      bool(QUOTA.search("agent 回報完成但找不到產出：crwv_2q26.json")), False)

print("\n== 財政季度分頁名（2026-08-12 SMCI/LITE 撞車事故的回歸測試）")
from quarters_cases import CASES  # noqa: E402  同目錄的資料表
from core.quarters import tab_name  # noqa: E402

for tk, fy_end, day, want in CASES:
    check(f"{tk} {day} → {want}",
          tab_name(tk, _dt.date.fromisoformat(day), fy_end=fy_end), want)

print("\n== sweep 的分頁季度推導（不能再從檔名猜）")
from sweep_podcasts import resolve_quarter  # noqa: E402


class _FakeTask:
    def __init__(self, d):
        self.call_at = _dt.datetime.fromisoformat(d)


class _FakeEdgar:
    def __init__(self, m):
        self.m = m

    def fiscal_year_end(self, tk):
        return self.m.get(tk)


_rows = {"SMCI": _FakeTask("2026-08-12T05:00:00"), "CRWV": _FakeTask("2026-08-12T05:00:00")}
_ed = _FakeEdgar({"SMCI": "0630", "CRWV": "1231"})
_T = "{ticker} {year}Q{q}"

check("非曆年公司：檔名的曆季會被更正",
      resolve_quarter("SMCI", "2026Q2", _rows, _ed, _T)[0], "2026Q4")
check("曆年公司：維持原樣",
      resolve_quarter("CRWV", "2026Q2", _rows, _ed, _T)[0], "2026Q2")
check("讀不到追蹤表就退回檔名，不能爆掉",
      resolve_quarter("SMCI", "2026Q2", {}, None, _T)[0], "2026Q2")
check("查不到 fiscalYearEnd 時退回曆季",
      resolve_quarter("SMCI", "2026Q2", _rows, _FakeEdgar({}), _T)[0], "2026Q2")

print("\n== .docx → Google Doc 分頁：內容不得有缺漏")
# 這一組是為了擋住 2026-08-12 那個缺陷：分頁只有 4,848 字、0 張表、0 張圖，
# 本機 memo 卻有 8,339 字、5 張表、1 張圖，而流程回報「成功」。
# 根因是 JSON 渲染器只認得 tldr 與 sections[].items，dashboard／table／qa 全掉了。
# 這裡直接拿 .docx 當基準，逐段逐格檢查 parse() 有沒有漏東西。
import docx as _docx

from core import docx_to_docs as _dd

_memo_dir = Path(__file__).resolve().parent.parent.parent / "MEMO" / "2Q26"
for _p in sorted(_memo_dir.glob("*.docx")):
    _raw = _docx.Document(str(_p))
    _blocks = _dd.parse(_p)
    _flat = _dd.normalize(_dd.plain_text(_blocks))

    check(f"{_p.stem}：表格數",
          sum(1 for b in _blocks if b["kind"] == "table"), len(_raw.tables))
    check(f"{_p.stem}：圖片數",
          sum(1 for b in _blocks if b["kind"] == "img"), len(_raw.inline_shapes))

    _lost = [par.text for par in _raw.paragraphs
             if par.text.strip() and _dd.normalize(par.text) not in _flat]
    check(f"{_p.stem}：段落無缺漏（少了 {len(_lost)} 段）", _lost[:2], [])

    _lost_cells = [c.text for t in _raw.tables for row in t.rows for c in row.cells
                   if c.text.strip() and _dd.normalize(c.text) not in _flat]
    check(f"{_p.stem}：表格格子無缺漏（少了 {len(_lost_cells)} 格）", _lost_cells[:2], [])

print("\n" + "=" * 50)
if failures:
    print(f"失敗 {len(failures)} 項：{failures}")
    raise SystemExit(1)
print("全部通過。")
