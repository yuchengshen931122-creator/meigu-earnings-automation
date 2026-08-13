#!/usr/bin/env python3
"""Render an earnings-memo JSON payload into the firm's standard .docx template.

Usage:
    python build_memo.py data.json "{專案根目錄}\\MEMO\\1Q26\\VZ 1Q26.docx"

The JSON schema is documented in references/schema.md. This script only
handles layout/formatting (fonts, bullets, table borders, page setup) --
all financial figures and Chinese-language content must already be correct
in the JSON before calling this script.
"""
import sys
import json
from docx import Document
from docx.shared import Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BODY_FONT = "Microsoft JhengHei"
HEADING_FONT = "Microsoft JhengHei"
PAGE_W_CM = 21.0
PAGE_H_CM = 29.7
MARGIN_CM = 2.0

# "Dark Red" (Word's standard swatch) rather than pure FF0000 -- reads as a deliberate highlight
# rather than an error/warning color, and holds up better in print.
RED = RGBColor(0xC0, 0x00, 0x00)

NAMED_COLORS = {"red": RED}


def resolve_color(color):
    if color is None:
        return None
    if color in NAMED_COLORS:
        return NAMED_COLORS[color]
    return RGBColor.from_string(color.lstrip("#").upper())


def set_font(run, name=BODY_FONT, east_asia=None, size=11, bold=False, color=None):
    east_asia = east_asia or name
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    rgb = resolve_color(color)
    if rgb is not None:
        run.font.color.rgb = rgb
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:cs"), name)
    rFonts.set(qn("w:eastAsia"), east_asia)


def add_run(paragraph, text, bold=False, size=11, font=BODY_FONT, color=None):
    run = paragraph.add_run(text)
    set_font(run, name=font, size=size, bold=bold, color=color)
    return run


def no_space(paragraph):
    """Zero out space-before/-after so body paragraphs sit flush against each other --
    python-docx's default template otherwise adds visible gaps between every bullet/line even
    without an extra blank line. Not applied to the title or section headings, which keep
    their own deliberate spacing for visual separation."""
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    return paragraph


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, text, bold=True, size=14)
    return p


def add_date(doc, text):
    p = doc.add_paragraph()
    no_space(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, text)
    return p


def add_heading(doc, text):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    add_run(p, text, bold=True, size=12, font=HEADING_FONT)
    return p


def configure_heading1_style(doc):
    style = doc.styles["Heading 1"]
    style.font.name = HEADING_FONT
    style.font.size = Pt(12)
    style.font.bold = True
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), HEADING_FONT)
    color_el = rPr.find(qn("w:color"))
    if color_el is not None:
        rPr.remove(color_el)


def render_item_segments(paragraph, item):
    """An item is either {"text", "bold", "color"} (single run) or
    {"segments": [{"text","bold","color"}, ...]} (mixed runs, same shape as tldr) -- lets a
    bullet highlight specific figures the same way tldr does, instead of only being plain or
    fully bold. "color": "red" marks the single most critical fact, on top of bold."""
    if "segments" in item:
        for seg in item["segments"]:
            add_run(paragraph, seg.get("text", ""), bold=bool(seg.get("bold", False)), color=seg.get("color"))
    else:
        add_run(paragraph, item.get("text", ""), bold=bool(item.get("bold", False)), color=item.get("color"))


# python-docx's built-in "List Bullet" / "List Bullet 2" styles render the same solid glyph at
# every level in its default template -- doesn't visually distinguish nesting the way the
# reference memos do (level 0 solid "●", level 1 hollow "○"). Define our own numbering instead
# of relying on those styles, so nesting is visually unambiguous regardless of which levels a
# given memo happens to use.
BULLET_CHARS = {0: "●", 1: "○", 2: "■"}  # ● ○ ■
BULLET_INDENTS = {0: (720, 360), 1: (1440, 360), 2: (2160, 360)}  # (left, hanging) in twips
_bullet_num_id = None


def ensure_bullet_numbering(doc):
    global _bullet_num_id
    if _bullet_num_id is not None:
        return _bullet_num_id

    root = doc.part.numbering_part.element
    abstract_num_id = 900
    num_id = 900

    abstractNum = OxmlElement("w:abstractNum")
    abstractNum.set(qn("w:abstractNumId"), str(abstract_num_id))
    for ilvl, char in BULLET_CHARS.items():
        left, hanging = BULLET_INDENTS[ilvl]
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(ilvl))
        numFmt = OxmlElement("w:numFmt")
        numFmt.set(qn("w:val"), "bullet")
        lvlText = OxmlElement("w:lvlText")
        lvlText.set(qn("w:val"), char)
        lvlJc = OxmlElement("w:lvlJc")
        lvlJc.set(qn("w:val"), "left")
        ind_pPr = OxmlElement("w:pPr")
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(left))
        ind.set(qn("w:hanging"), str(hanging))
        ind_pPr.append(ind)
        rPr = OxmlElement("w:rPr")
        rFonts = OxmlElement("w:rFonts")
        for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
            rFonts.set(qn(attr), BODY_FONT)
        rPr.append(rFonts)
        for child in (numFmt, lvlText, lvlJc, ind_pPr, rPr):
            lvl.append(child)
        abstractNum.append(lvl)

    # schema requires all <w:abstractNum> elements before any <w:num> element
    first_num_index = next(
        (i for i, child in enumerate(root) if child.tag == qn("w:num")), None
    )
    if first_num_index is None:
        root.append(abstractNum)
    else:
        root.insert(first_num_index, abstractNum)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstractNumId_ref = OxmlElement("w:abstractNumId")
    abstractNumId_ref.set(qn("w:val"), str(abstract_num_id))
    num.append(abstractNumId_ref)
    root.append(num)

    _bullet_num_id = num_id
    return num_id


def add_bullet(doc, item, level=0):
    # Element order inside <w:pPr> is schema-fixed (numPr, then spacing, then ind) -- build
    # numPr first, call no_space() (which inserts <w:spacing> correctly relative to whatever's
    # already there), then append <w:ind> last, so the final order comes out numPr/spacing/ind
    # rather than corrupting the paragraph properties sequence.
    num_id = ensure_bullet_numbering(doc)
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()

    numPr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), str(level))
    numId_el = OxmlElement("w:numId")
    numId_el.set(qn("w:val"), str(num_id))
    numPr.append(ilvl_el)
    numPr.append(numId_el)
    pPr.append(numPr)

    no_space(p)

    left, hanging = BULLET_INDENTS[level]
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), str(left))
    ind.set(qn("w:hanging"), str(hanging))
    pPr.append(ind)

    render_item_segments(p, item)
    return p


def add_numbered(doc, segments):
    p = doc.add_paragraph(style="List Number")
    no_space(p)
    for seg in segments:
        add_run(p, seg.get("text", ""), bold=bool(seg.get("bold", False)), color=seg.get("color"))
    return p


def set_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "8")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    tblPr.append(borders)


def set_table_width(table, width_emu):
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    twips = int(width_emu / 635)  # 1 twip = 635 EMU
    tblW.set(qn("w:w"), str(twips))
    tblW.set(qn("w:type"), "dxa")


def render_cell(cell, value, default_bold=False):
    """A cell is a plain string, or {"text","bold","color"}, or a list of those segments --
    the same shape used everywhere else, so a table can bold/redden one figure inside a cell
    instead of the whole column being uniform."""
    p = cell.paragraphs[0]
    no_space(p)
    if isinstance(value, list):
        for seg in value:
            add_run(p, seg.get("text", ""), bold=bool(seg.get("bold", False)), color=seg.get("color"))
    elif isinstance(value, dict):
        add_run(p, value.get("text", ""), bold=bool(value.get("bold", default_bold)), color=value.get("color"))
    else:
        add_run(p, str(value), bold=default_bold)


def column_ratios(n, col_ratio=None):
    """Explicit ratios win (normalised so they always sum to 1). Otherwise the first column
    holds the row label and gets the most room, and the rest split what's left evenly.

    The old default of (0.34, 0.66) summed to 1.0, which left nothing for a third column and
    silently rendered every column past the second at zero width -- any table wider than two
    columns came out broken. Never reintroduce a default whose entries already sum to 1."""
    if col_ratio:
        if len(col_ratio) != n:
            raise ValueError(f"col_ratio has {len(col_ratio)} entries but the table has {n} columns")
        total = float(sum(col_ratio))
        return [r / total for r in col_ratio]
    if n == 1:
        return [1.0]
    if n == 2:
        return [0.34, 0.66]
    first = 0.22 if n >= 5 else 0.28
    return [first] + [(1.0 - first) / (n - 1)] * (n - 1)


def add_table(doc, columns, rows, col_ratio=None):
    table = doc.add_table(rows=1, cols=len(columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)

    hdr_cells = table.rows[0].cells
    for i, col in enumerate(columns):
        no_space(hdr_cells[i].paragraphs[0])
        add_run(hdr_cells[i].paragraphs[0], str(col), bold=True)

    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            render_cell(cells[i], val, default_bold=(i == 0))

    content_width_cm = PAGE_W_CM - 2 * MARGIN_CM
    ratios = column_ratios(len(columns), col_ratio)
    widths = [Emu(Cm(content_width_cm * r)) for r in ratios]
    for i, width in enumerate(widths):
        table.columns[i].width = width
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = widths[i]
    set_table_width(table, sum(widths))
    return table


def render_prose_segments(paragraph, item):
    """Like render_item_segments, but for a bare string field (qa["q"]/["a"]) that may
    instead be given as a segments array under a "_segments" sibling key."""
    if isinstance(item, list):
        for seg in item:
            add_run(paragraph, seg.get("text", ""), bold=bool(seg.get("bold", False)), color=seg.get("color"))
    else:
        add_run(paragraph, item)


def add_image(doc, image_path, caption=None, max_width_cm=None, max_height_cm=20.0):
    from PIL import Image as PILImage

    content_width_cm = PAGE_W_CM - 2 * MARGIN_CM
    max_width_cm = max_width_cm or content_width_cm

    with PILImage.open(image_path) as im:
        px_w, px_h = im.size
    aspect = px_h / px_w

    width_cm = max_width_cm
    height_cm = width_cm * aspect
    if height_cm > max_height_cm:
        height_cm = max_height_cm
        width_cm = height_cm / aspect

    p = doc.add_paragraph()
    no_space(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(image_path, width=Cm(width_cm), height=Cm(height_cm))

    if caption:
        cap_p = doc.add_paragraph()
        no_space(cap_p)
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(cap_p, caption, size=9)
    return p


def add_block_caption(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    add_run(p, text, bold=True, size=11)
    return p


def add_dashboard(doc, blocks):
    """The page-1 scan block: a short stack of captioned tables rendered between the tldr and
    the income-statement screenshot. Numbers live here; the tldr carries what numbers can't."""
    for blk in blocks:
        if blk.get("heading"):
            add_block_caption(doc, blk["heading"])
        add_table(doc, blk["columns"], blk["rows"], col_ratio=blk.get("col_ratio"))
        if blk.get("note"):
            p = doc.add_paragraph()
            no_space(p)
            add_run(p, blk["note"], size=9)


def add_qa_section(doc, section_data):
    for theme in section_data.get("themes", []):
        add_bullet(doc, theme, level=0)
    for i, qa in enumerate(section_data.get("qa", []), start=1):
        q_p = doc.add_paragraph()
        no_space(q_p)
        add_run(q_p, f"Q{i}：", bold=True)
        render_prose_segments(q_p, qa.get("q_segments", qa.get("q", "")))

        a_p = doc.add_paragraph()
        no_space(a_p)
        add_run(a_p, "A：", bold=True)
        render_prose_segments(a_p, qa.get("a_segments", qa.get("a", "")))


def build(data, output_path):
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(PAGE_W_CM)
    section.page_height = Cm(PAGE_H_CM)
    section.left_margin = Cm(MARGIN_CM)
    section.right_margin = Cm(MARGIN_CM)
    section.top_margin = Cm(MARGIN_CM)
    section.bottom_margin = Cm(MARGIN_CM)

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal_rPr = normal.element.get_or_add_rPr()
    normal_rFonts = normal_rPr.find(qn("w:rFonts"))
    if normal_rFonts is None:
        normal_rFonts = OxmlElement("w:rFonts")
        normal_rPr.append(normal_rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        normal_rFonts.set(qn(attr), BODY_FONT)

    configure_heading1_style(doc)

    add_title(doc, data["title"])
    add_date(doc, data["release_date"])

    for bullet_segments in data["tldr"]:
        add_numbered(doc, bullet_segments)

    dashboard = data.get("dashboard")
    if dashboard:
        add_dashboard(doc, dashboard)

    income_statement_image = data.get("income_statement_image")
    if income_statement_image:
        add_image(
            doc,
            income_statement_image["path"],
            caption=income_statement_image.get("caption"),
        )

    for section_data in data["sections"]:
        add_heading(doc, section_data["heading"])
        stype = section_data["type"]

        if stype == "bullets":
            for item in section_data["items"]:
                add_bullet(doc, item, level=0)
                for child in item.get("children", []):
                    add_bullet(doc, child, level=1)

        elif stype == "table":
            add_table(doc, section_data["columns"], section_data["rows"], col_ratio=section_data.get("col_ratio"))

        elif stype == "qa":
            add_qa_section(doc, section_data)

        else:
            raise ValueError(f"Unknown section type: {stype!r}")

    doc.save(output_path)


def main():
    if len(sys.argv) != 3:
        print(f"Usage: python {sys.argv[0]} <data.json> <output.docx>", file=sys.stderr)
        sys.exit(1)
    data_path, output_path = sys.argv[1], sys.argv[2]
    with open(data_path, encoding="utf-8") as f:
        data = json.load(f)
    build(data, output_path)
    print(f"Wrote {output_path}")


if __name__ == "__main__":
    main()
