"""Insert an income-statement PNG + caption into an EXISTING memo .docx without touching any text.

Places the image exactly where build_memo.py would: after the TL;DR numbered list, immediately
before the "一、" heading. Matches build_memo.py's add_image(): centred, 17cm content width,
20cm max height, 9pt caption in Microsoft JhengHei.

Idempotent: if the target docx already contains an inline image, it refuses and exits 2.

usage: python insert_is.py <docx> <png> "<caption>"
"""
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image as PILImage

BODY_FONT = "Microsoft JhengHei"
PAGE_W_CM = 21.0
MARGIN_CM = 2.0
MAX_H_CM = 20.0


def no_space(p):
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)


def set_font(run, size=11):
    run.font.name = BODY_FONT
    run.font.size = Pt(size)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(attr), BODY_FONT)


def main(docx_path, png_path, caption):
    doc = Document(docx_path)

    if doc.inline_shapes:
        print(f"SKIP {docx_path}: already has {len(doc.inline_shapes)} inline image(s)")
        return 2

    target = None
    for p in doc.paragraphs:
        if p.text.strip().startswith("一、"):
            target = p
            break
    if target is None:
        print(f"FAIL {docx_path}: no paragraph starting with 一、")
        return 1

    content_w = PAGE_W_CM - 2 * MARGIN_CM
    with PILImage.open(png_path) as im:
        px_w, px_h = im.size
    aspect = px_h / px_w
    w = content_w
    h = w * aspect
    if h > MAX_H_CM:
        h = MAX_H_CM
        w = h / aspect

    img_p = target.insert_paragraph_before()
    no_space(img_p)
    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_p.add_run().add_picture(png_path, width=Cm(w), height=Cm(h))

    cap_p = target.insert_paragraph_before()
    no_space(cap_p)
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(cap_p.add_run(caption), size=9)

    doc.save(docx_path)
    print(f"OK {docx_path}: image {w:.1f}x{h:.1f}cm inserted before 一、")
    return 0


if __name__ == "__main__":
    sys.exit(main(sys.argv[1], sys.argv[2], sys.argv[3]))
